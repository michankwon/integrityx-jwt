"""
Document Intelligence Service for Enhanced Information Extraction

This module provides AI-powered document processing capabilities including:
- Automatic document type classification
- OCR text extraction from images and PDFs
- Structured data extraction from various document formats
- Smart form auto-population
- Business rule validation
"""

import json
import re
import logging
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
from pathlib import Path
import hashlib

# Document processing imports
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

logger = logging.getLogger(__name__)


class DocumentIntelligenceService:
    """
    AI-powered document intelligence service for extracting structured data
    from various document types and formats.
    """
    
    def __init__(self):
        """Initialize the document intelligence service."""
        self.document_classifiers = {
            'loan_application': {
                'keywords': ['application', 'borrower', 'income', 'employment', 'loan request'],
                'patterns': [r'loan\s+application', r'borrower\s+information', r'income\s+verification']
            },
            'credit_report': {
                'keywords': ['credit', 'score', 'bureau', 'fico', 'transunion', 'equifax', 'experian'],
                'patterns': [r'credit\s+report', r'fico\s+score', r'credit\s+bureau']
            },
            'appraisal': {
                'keywords': ['appraisal', 'property', 'value', 'appraiser', 'market value'],
                'patterns': [r'appraisal\s+report', r'property\s+value', r'market\s+analysis']
            },
            'underwriting': {
                'keywords': ['underwriting', 'approval', 'conditions', 'risk assessment'],
                'patterns': [r'underwriting\s+decision', r'loan\s+approval', r'risk\s+assessment']
            },
            'closing_documents': {
                'keywords': ['closing', 'settlement', 'hud', 'disclosure', 'final'],
                'patterns': [r'closing\s+disclosure', r'settlement\s+statement', r'hud\s+form']
            },
            'income_verification': {
                'keywords': ['income', 'payroll', 'w2', 'tax return', 'employment'],
                'patterns': [r'income\s+verification', r'payroll\s+stub', r'w-2\s+form']
            }
        }
        
        self.data_extractors = {
            'loan_id': [r'loan\s*id[:\s]*([A-Za-z0-9_-]+)', r'loan\s*number[:\s]*([A-Za-z0-9_-]+)'],
            'borrower_name': [r'borrower[:\s]*([A-Za-z\s]+)', r'name[:\s]*([A-Za-z\s]+)'],
            'property_address': [r'property\s*address[:\s]*([^\n]+)', r'address[:\s]*([^\n]+)'],
            'loan_amount': [r'loan\s*amount[:\s]*\$?([0-9,]+)', r'amount[:\s]*\$?([0-9,]+)'],
            'interest_rate': [r'interest\s*rate[:\s]*([0-9.]+)%?', r'rate[:\s]*([0-9.]+)%?'],
            'loan_term': [r'term[:\s]*([0-9]+)\s*(months?|years?)', r'loan\s*term[:\s]*([0-9]+)']
        }
    
    def extract_structured_data(self, file_content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """
        Extract structured data from any document type.
        
        Args:
            file_content: Raw file content as bytes
            filename: Original filename
            content_type: MIME content type
            
        Returns:
            Dict containing extracted structured data
        """
        try:
            # Determine file type and extract accordingly
            file_extension = Path(filename).suffix.lower()
            
            if content_type == 'application/json' or file_extension == '.json':
                return self._extract_from_json(file_content)
            elif content_type == 'application/pdf' or file_extension == '.pdf':
                return self._extract_from_pdf(file_content)
            elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                                 'application/msword'] or file_extension in ['.docx', '.doc']:
                return self._extract_from_word(file_content)
            elif content_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                                 'application/vnd.ms-excel'] or file_extension in ['.xlsx', '.xls']:
                return self._extract_from_excel(file_content)
            elif content_type.startswith('image/') or file_extension in ['.jpg', '.jpeg', '.png', '.tiff']:
                return self._extract_from_image(file_content)
            elif content_type == 'text/plain' or file_extension == '.txt':
                return self._extract_from_text(file_content)
            else:
                return self._extract_basic_info(file_content, filename, content_type)
                
        except Exception as e:
            logger.error(f"Error extracting structured data: {e}")
            return self._extract_basic_info(file_content, filename, content_type)
    
    def _extract_from_json(self, content: bytes) -> Dict[str, Any]:
        """Extract data from JSON documents."""
        try:
            json_data = json.loads(content.decode('utf-8'))
            
            # Extract common loan fields
            extracted = {
                'document_type': 'json',
                'raw_data': json_data,
                'extracted_fields': {}
            }
            
            # Extract specific fields
            for field, patterns in self.data_extractors.items():
                value = self._extract_field_from_data(json_data, field, patterns)
                if value:
                    extracted['extracted_fields'][field] = value
            
            # Classify document type
            extracted['document_classification'] = self._classify_document(json_data)
            
            return extracted
            
        except Exception as e:
            logger.error(f"Error processing JSON: {e}")
            return {'document_type': 'json', 'error': str(e)}
    
    def _extract_from_pdf(self, content: bytes) -> Dict[str, Any]:
        """Extract data from PDF documents."""
        if not HAS_PYPDF2:
            return {'document_type': 'pdf', 'error': 'PyPDF2 not available'}
        
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text_content = ""
            
            # Extract text from all pages
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
            
            extracted = {
                'document_type': 'pdf',
                'text_content': text_content,
                'page_count': len(pdf_reader.pages),
                'extracted_fields': {}
            }
            
            # Extract structured data from text
            for field, patterns in self.data_extractors.items():
                value = self._extract_field_from_text(text_content, patterns)
                if value:
                    extracted['extracted_fields'][field] = value
            
            # Classify document type
            extracted['document_classification'] = self._classify_document(text_content)
            
            return extracted
            
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            return {'document_type': 'pdf', 'error': str(e)}
    
    def _extract_from_word(self, content: bytes) -> Dict[str, Any]:
        """Extract data from Word documents."""
        if not HAS_DOCX:
            return {'document_type': 'word', 'error': 'python-docx not available'}
        
        try:
            doc = DocxDocument(io.BytesIO(content))
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            extracted = {
                'document_type': 'word',
                'text_content': text_content,
                'extracted_fields': {}
            }
            
            # Extract structured data from text
            for field, patterns in self.data_extractors.items():
                value = self._extract_field_from_text(text_content, patterns)
                if value:
                    extracted['extracted_fields'][field] = value
            
            # Classify document type
            extracted['document_classification'] = self._classify_document(text_content)
            
            return extracted
            
        except Exception as e:
            logger.error(f"Error processing Word document: {e}")
            return {'document_type': 'word', 'error': str(e)}
    
    def _extract_from_excel(self, content: bytes) -> Dict[str, Any]:
        """Extract data from Excel documents."""
        if not HAS_OPENPYXL:
            return {'document_type': 'excel', 'error': 'openpyxl not available'}
        
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(content))
            extracted = {
                'document_type': 'excel',
                'worksheets': [],
                'extracted_fields': {}
            }
            
            # Process each worksheet
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = []
                
                # Extract data from cells
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        sheet_data.append([str(cell) if cell is not None else "" for cell in row])
                
                extracted['worksheets'].append({
                    'name': sheet_name,
                    'data': sheet_data
                })
            
            # Try to extract structured data from the first sheet
            if extracted['worksheets']:
                first_sheet_data = extracted['worksheets'][0]['data']
                text_content = " ".join([" ".join(row) for row in first_sheet_data])
                
                for field, patterns in self.data_extractors.items():
                    value = self._extract_field_from_text(text_content, patterns)
                    if value:
                        extracted['extracted_fields'][field] = value
            
            return extracted
            
        except Exception as e:
            logger.error(f"Error processing Excel document: {e}")
            return {'document_type': 'excel', 'error': str(e)}
    
    def _extract_from_image(self, content: bytes) -> Dict[str, Any]:
        """Extract data from image documents using OCR."""
        if not HAS_OCR:
            return {'document_type': 'image', 'error': 'OCR not available'}
        
        try:
            image = Image.open(io.BytesIO(content))
            text_content = pytesseract.image_to_string(image)
            
            extracted = {
                'document_type': 'image',
                'text_content': text_content,
                'extracted_fields': {}
            }
            
            # Extract structured data from OCR text
            for field, patterns in self.data_extractors.items():
                value = self._extract_field_from_text(text_content, patterns)
                if value:
                    extracted['extracted_fields'][field] = value
            
            # Classify document type
            extracted['document_classification'] = self._classify_document(text_content)
            
            return extracted
            
        except Exception as e:
            logger.error(f"Error processing image with OCR: {e}")
            return {'document_type': 'image', 'error': str(e)}
    
    def _extract_from_text(self, content: bytes) -> Dict[str, Any]:
        """Extract data from plain text documents."""
        try:
            text_content = content.decode('utf-8')
            
            extracted = {
                'document_type': 'text',
                'text_content': text_content,
                'extracted_fields': {}
            }
            
            # Extract structured data from text
            for field, patterns in self.data_extractors.items():
                value = self._extract_field_from_text(text_content, patterns)
                if value:
                    extracted['extracted_fields'][field] = value
            
            # Classify document type
            extracted['document_classification'] = self._classify_document(text_content)
            
            return extracted
            
        except Exception as e:
            logger.error(f"Error processing text document: {e}")
            return {'document_type': 'text', 'error': str(e)}
    
    def _extract_basic_info(self, content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """Extract basic information when specific extraction fails."""
        return {
            'document_type': 'unknown',
            'filename': filename,
            'content_type': content_type,
            'file_size': len(content),
            'extracted_fields': {},
            'document_classification': 'unknown'
        }
    
    def _extract_field_from_text(self, text: str, patterns: List[str]) -> Optional[str]:
        """Extract field value from text using regex patterns."""
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return None
    
    def _extract_field_from_data(self, data: Dict[str, Any], field: str, patterns: List[str]) -> Optional[str]:
        """Extract field value from structured data."""
        # First try direct field access
        if field in data:
            return str(data[field])
        
        # Then try nested field access
        if '.' in field:
            keys = field.split('.')
            current = data
            for key in keys:
                if isinstance(current, dict) and key in current:
                    current = current[key]
                else:
                    return None
            return str(current)
        
        # Finally try pattern matching on string representation
        text = json.dumps(data)
        return self._extract_field_from_text(text, patterns)
    
    def _classify_document(self, content: str) -> str:
        """Classify document type based on content analysis."""
        content_lower = content.lower()
        
        best_match = 'unknown'
        best_score = 0
        
        for doc_type, classifier in self.document_classifiers.items():
            score = 0
            
            # Check keywords
            for keyword in classifier['keywords']:
                if keyword in content_lower:
                    score += 1
            
            # Check patterns
            for pattern in classifier['patterns']:
                if re.search(pattern, content_lower):
                    score += 2
            
            if score > best_score:
                best_score = score
                best_match = doc_type
        
        return best_match if best_score > 0 else 'unknown'
    
    def auto_populate_form(self, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
        """Auto-populate form fields from extracted data."""
        form_data = {}
        
        # Map extracted fields to form fields
        field_mapping = {
            'loan_id': 'loanId',
            'borrower_name': 'borrowerName',
            'property_address': 'propertyAddress',
            'loan_amount': 'amount',
            'interest_rate': 'rate',
            'loan_term': 'term'
        }
        
        extracted_fields = extracted_data.get('extracted_fields', {})
        
        for extracted_field, form_field in field_mapping.items():
            if extracted_field in extracted_fields:
                form_data[form_field] = extracted_fields[extracted_field]
        
        # Add document classification
        form_data['documentType'] = extracted_data.get('document_classification', 'unknown')
        
        return form_data
    
    def validate_business_rules(self, extracted_data: Dict[str, Any]) -> Tuple[bool, List[str]]:
        """Validate extracted data against business rules."""
        errors = []
        extracted_fields = extracted_data.get('extracted_fields', {})
        
        # Validate loan amount
        if 'loan_amount' in extracted_fields:
            try:
                amount = float(extracted_fields['loan_amount'].replace(',', ''))
                if amount <= 0:
                    errors.append("Loan amount must be positive")
                elif amount > 10000000:  # 10 million max
                    errors.append("Loan amount exceeds maximum limit")
            except ValueError:
                errors.append("Invalid loan amount format")
        
        # Validate interest rate
        if 'interest_rate' in extracted_fields:
            try:
                rate = float(extracted_fields['interest_rate'])
                if rate < 0 or rate > 50:  # 50% max
                    errors.append("Interest rate must be between 0 and 50%")
            except ValueError:
                errors.append("Invalid interest rate format")
        
        # Validate loan term
        if 'loan_term' in extracted_fields:
            try:
                term = int(extracted_fields['loan_term'])
                if term <= 0 or term > 3600:  # 30 years max
                    errors.append("Loan term must be between 1 and 3600 months")
            except ValueError:
                errors.append("Invalid loan term format")
        
        return len(errors) == 0, errors


# Example usage and testing
if __name__ == "__main__":
    # Test the document intelligence service
    service = DocumentIntelligenceService()
    
    # Test with sample JSON data
    sample_json = {
        "loanId": "LOAN_001",
        "amount": 250000,
        "rate": 6.5,
        "term": 360,
        "borrower": {
            "name": "John Smith",
            "email": "john@example.com"
        }
    }
    
    result = service._extract_from_json(json.dumps(sample_json).encode())
    print("Extraction result:", json.dumps(result, indent=2))
    
    # Test form auto-population
    form_data = service.auto_populate_form(result)
    print("Form data:", json.dumps(form_data, indent=2))
    
    # Test business rule validation
    is_valid, errors = service.validate_business_rules(result)
    print("Validation result:", is_valid, errors)
